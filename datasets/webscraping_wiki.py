import requests
from bs4 import BeautifulSoup
from docx import Document

# Wikipedia page URL
url = "https://en.wikipedia.org/wiki/Physical_fitness"  # Change to any Wikipedia page

# Set headers to avoid blocking
headers = {"User-Agent": "Mozilla/5.0"}
response = requests.get(url, headers=headers)

# Parse the page content
soup = BeautifulSoup(response.text, "html.parser")


paragraphs = soup.find_all("p")
lists = soup.find_all("ul")

# Create a Word document
doc = Document()
# doc.add_heading("Wikipedia Article: Apple", level=1)  # Change title accordingly

# Add <p> tag content
for para in paragraphs:
    text = para.get_text(strip=True)
    if text:  # Skip empty paragraphs
        doc.add_paragraph(text)

# Add <ul> list content as paragraphs (no bullets)
selected_lists = lists[41:48]  # Extract only ul[41] to ul[47]
for ul in lists:
    list_items = ul.find_all("li")  # Extract <li> items inside <ul>
    for li in list_items:
        text = li.get_text(strip=True)
        if text:
            doc.add_paragraph(text)  # Add list items as normal paragraphs

# Save to a .docx file
# i have changed file name so as to not tamper the real one
doc.save("wikipedia_fitness_article2.docx")

print("Wikipedia content saved as 'wikipedia_fitness_article2.docx'!")